#!/usr/bin/env python3
"""Создание документа Токеномика в формате Word."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Заливка ячейки."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'
    elif level == 2:
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    elif level == 3:
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    run.font.italic = True
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc, headers, rows, header_bg='D9E2F3'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
        set_cell_shading(cell, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

def add_formula(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Consolas'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    return p

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Титул
add_heading(doc, 'ТОКЕНОМИКА ПРОТОКОЛА «ЕЛЕНА» (Elena Protocol)', 0)
add_heading(doc, 'Детальное описание экономической модели', 2)
p = doc.add_paragraph()
r = p.add_run('Версия 2.0 | Март 2026')
r.font.size = Pt(12)
r.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph('_' * 80)

# Содержание
add_heading(doc, 'Содержание', 1)
contents = [
    '1. Базовые принципы',
    '2. Единицы измерения',
    '3. Эмиссия: механизм «Осаждения» (3.1–3.4)',
    '4. Комиссии (4.1–4.4)',
    '5. Репутационная система (5.1–5.2)',
    '6. Стейкинг репутации (6.1–6.2)',
    '7. Дефляционные механизмы (7.1–7.3)',
    '8. Распределение начальной эмиссии',
    '9. Сравнение с другими сетями',
    '10. Математическое резюме',
    '11. Прогнозы и моделирование',
    '12. Управление параметрами (Governance)',
]
for c in contents:
    add_para(doc, c)

doc.add_page_break()

# 1
add_heading(doc, '1. Базовые принципы', 1)
add_para(doc, 'Экономика «Елены» построена на трёх фундаментальных принципах:')
add_table(doc, ['Принцип', 'Суть'], [
    ['Репутация, а не богатство', 'Власть в сети определяется вкладом в её развитие, а не количеством монет'],
    ['Естественная эмиссия', 'Новые монеты не майнятся и не стейкаются, а «выпадают в осадок» пропорционально полезной работе'],
    ['Саморегуляция', 'Параметры экономики меняются через голосование, где вес голоса = репутация'],
])

# 2
add_heading(doc, '2. Единицы измерения', 1)
add_para(doc, 'Для удобства и точности расчётов используются две единицы:')
add_table(doc, ['Единица', 'Значение', 'Использование'], [
    ['ELENA', '1 ELENA', 'Отображение в интерфейсах, кошельках'],
    ['микро-ELENA', '1 ELENA = 1 000 000 микро', 'Все внутренние расчёты, комиссии, эмиссия'],
])
add_code(doc, 'pub const MICRO_PER_ELENA: u64 = 1_000_000;')

# 3
add_heading(doc, '3. Эмиссия: механизм «Осаждения»', 1)

add_heading(doc, '3.1 Базовая формула', 2)
add_para(doc, 'Новые монеты «выпадают в осадок» подобно тому, как кристаллы растут в перенасыщенном растворе:')
add_formula(doc, 'награда(узел) = база_эмиссии × (байты × время) / общее_хранилище × фактор_репутации')
add_para(doc, 'Где:')
add_para(doc, '• байты = объём данных транзакций, которые узел хранит (в байтах)', bold=False)
add_para(doc, '• время = как долго узел хранит эти данные (в часах)', bold=False)
add_para(doc, '• общее_хранилище = сумма (байты × время) по всем узлам сети', bold=False)
add_para(doc, '• фактор_репутации = множитель от 0.5 до 2.0 на основе репутации узла', bold=False)

add_heading(doc, '3.2 Параметры эмиссии', 2)
add_code(doc, 'pub const MAX_SUPPLY_ELENA: u64 = 21_000_000;           // Максимальное предложение')
add_code(doc, 'pub const MAX_SUPPLY_MICRO: u64 = 21_000_000 * 1_000_000;')
add_code(doc, 'pub const EMISSION_BASE_PER_HOUR_MICRO: u64 = 1_000_000; // 1 ELENA в час за 1 МБ·день')
add_code(doc, 'pub const EMISSION_HALVING_INTERVAL_YEARS: u64 = 2;      // Халвинг каждые 2 года')
add_code(doc, 'pub const EMISSION_START_YEAR: u64 = 2026;               // Год запуска')

add_heading(doc, '3.3 График эмиссии', 2)
add_table(doc, ['Период', 'Год', 'Годовая эмиссия', 'Накоплено (приблизительно)'], [
    ['Запуск', '2026', '1 000 000', '1 000 000'],
    ['Халвинг 1', '2028', '500 000', '2 500 000'],
    ['Халвинг 2', '2030', '250 000', '3 500 000'],
    ['Халвинг 3', '2032', '125 000', '4 000 000'],
    ['...', '...', '...', '...'],
    ['Итого', '~2036', '~21 000 000', '21 000 000'],
])
add_para(doc, 'График повторяет биткоиновский принцип убывающей эмиссии, но с совершенно другим механизмом распределения.')

add_heading(doc, '3.4 Примеры расчёта', 2)
add_para(doc, 'Пример 1: Средний узел')
add_para(doc, '• Хранит: 100 МБ транзакций')
add_para(doc, '• Время: 24 часа')
add_para(doc, '• Репутация: 0.7')
add_para(doc, '• Общее хранилище сети: 10 000 МБ·день')
add_formula(doc, 'награда = 1 ELENA × (100 × 24) / 10 000 × (0.5 + 0.5 × 0.7)')
add_formula(doc, '        = 1 × 2400 / 10 000 × 0.85')
add_formula(doc, '        = 0.204 ELENA за день')
add_para(doc, 'Пример 2: Крупный узел с высокой репутацией')
add_para(doc, '• Хранит: 1 000 МБ')
add_para(doc, '• Время: 24 часа')
add_para(doc, '• Репутация: 0.95')
add_para(doc, '• Общее хранилище: 10 000 МБ·день')
add_formula(doc, 'награда = 1 × (1000 × 24) / 10 000 × (0.5 + 0.5 × 0.95)')
add_formula(doc, '        = 1 × 24 000 / 10 000 × 0.975')
add_formula(doc, '        = 2.34 ELENA за день')

# 4
add_heading(doc, '4. Комиссии', 1)

add_heading(doc, '4.1 Формула комиссии', 2)
add_formula(doc, 'комиссия(микро) = FEE_BASE_MICRO + (amount_микро × FEE_RATE_BP / 10_000) × priority_mult')
add_code(doc, 'pub const FEE_BASE_MICRO: u64 = 100;   // 0.0001 ELENA')
add_code(doc, 'pub const FEE_RATE_BP: u64 = 1;        // 0.01% (1 базисный пункт)')

add_heading(doc, '4.2 Приоритеты транзакций', 2)
add_table(doc, ['Приоритет', 'Множитель', 'Когда использовать'], [
    ['Normal', '×1', 'Обычные переводы'],
    ['Urgent', '×2', 'Срочные платежи'],
    ['Critical', '×10', 'Крупные суммы, требующие быстрого подтверждения'],
])

add_heading(doc, '4.3 Бесплатные микротранзакции', 2)
add_para(doc, 'Для стимулирования микроплатежей действует правило:')
add_formula(doc, 'if amount < 0.01 ELENA (10 000 микро) AND reputation ≥ 0.8: fee = 0')

add_heading(doc, '4.4 Распределение комиссий', 2)
add_code(doc, 'pub const FEE_SHARE_STORAGE: u64 = 50;  // 50% - узлам, хранящим транзакцию')
add_code(doc, 'pub const FEE_SHARE_RELAY: u64 = 30;    // 30% - узлам-ретрансляторам')
add_code(doc, 'pub const FEE_SHARE_BURN: u64 = 20;     // 20% - сжигается')
add_para(doc, 'Механизм распределения:')
add_para(doc, '1. 50% получают узлы, в чьих локальных графах есть эта транзакция (пропорционально времени хранения)')
add_para(doc, '2. 30% распределяется между узлами, участвовавшими в распространении')
add_para(doc, '3. 20% отправляется на burn-адрес (0x0000...), навсегда удаляясь из оборота')
add_para(doc, 'Пример: Транзакция с комиссией 0.01 ELENA')
add_para(doc, '• Хранители получают: 0.005 ELENA (суммарно)')
add_para(doc, '• Ретрансляторы: 0.003 ELENA')
add_para(doc, '• Сжигается: 0.002 ELENA')

# 5
add_heading(doc, '5. Репутационная система', 1)

add_heading(doc, '5.1 Базовые правила', 2)
add_para(doc, 'Репутация измеряется в диапазоне 0.01 – 0.99 и обновляется по формуле:')
add_formula(doc, 'R_нов = min(0.99, max(0.01, R_стар + Δ))')

add_heading(doc, '5.2 Дельта-таблица', 2)
add_table(doc, ['Действие', 'Δ репутации', 'Примечание'], [
    ['Пересылка валидной транзакции', '+0.0005', 'За каждую пересланную'],
    ['Создание Alert\'а', '+0.01', 'За обнаружение коллизии'],
    ['Хранение данных (за день)', '+0.001', 'Автоматически'],
    ['Ежедневное затухание', '-0.0001', 'Если узел неактивен'],
    ['Двойная трата', 'сброс до 0.01', 'Единоразово'],
    ['Сговор с атакующим', '-0.5 × R', 'Пропорционально'],
])
add_code(doc, 'pub const REPUTATION_PUNISH_MIN: f64 = 0.01;')
add_code(doc, 'pub const REPUTATION_DELTA_RELAY: f64 = 0.0005;')
add_code(doc, 'pub const REPUTATION_DELTA_ALERT: f64 = 0.01;')
add_code(doc, 'pub const REPUTATION_DELTA_STORAGE_PER_DAY: f64 = 0.001;')
add_code(doc, 'pub const REPUTATION_DECAY_PER_DAY: f64 = 0.0001;')

# 6
add_heading(doc, '6. Стейкинг репутации', 1)

add_heading(doc, '6.1 Механизм', 2)
add_para(doc, 'Узлы могут заморозить часть своей репутации (временно снизить её), чтобы получить повышенное влияние:')
add_formula(doc, 'r_эфф = r × (1 + 0.5 × s), где s ∈ [0, 0.5]')
add_para(doc, '• s = доля замороженной репутации (от 0 до 50%)')
add_para(doc, '• r_эфф используется для: множителя эмиссии; вероятности быть выбранным арбитром; веса в голосованиях')

add_heading(doc, '6.2 Риски', 2)
add_para(doc, 'Если узел с замороженной репутацией совершает нарушение:')
add_formula(doc, 'штраф = r × (1 + s)  // Потеря репутации увеличивается пропорционально стейку')
add_formula(doc, 'замороженная_часть = 0  // Полностью сгорает')

# 7
add_heading(doc, '7. Дефляционные механизмы', 1)

add_heading(doc, '7.1 Сжигание комиссий', 2)
add_para(doc, '20% всех комиссий навсегда удаляются из оборота. Это создаёт:')
add_para(doc, '• Естественное дефляционное давление')
add_para(doc, '• Компенсацию инфляции от эмиссии')
add_para(doc, '• Рост ценности при увеличении использования сети')

add_heading(doc, '7.2 Штрафы за нарушения', 2)
add_para(doc, 'При подтверждённой попытке двойной траты:')
add_para(doc, '1. Репутация падает до 0.01')
add_para(doc, '2. 1% от баланса сжигается (если баланс > 100 ELENA)')

add_heading(doc, '7.3 Затухание неактивных счетов', 2)
add_para(doc, 'Если адрес не использовался более 1 года:')
add_para(doc, '• Баланс уменьшается на 1% в месяц')
add_para(doc, '• Репутация падает до 0.01')
add_para(doc, '• Счёт переходит в «спящий» режим')

# 8
add_heading(doc, '8. Распределение начальной эмиссии', 1)
add_para(doc, 'Принцип: АБСОЛЮТНАЯ СПРАВЕДЛИВОСТЬ', bold=True)
add_table(doc, ['Механизм', 'Биткоин', 'Елена'], [
    ['Genesis block', '50 BTC (Сатоши мог майнить в одиночку)', '0 ELENA'],
    ['Пре-майн', '0 (но ранние майнеры получили всё)', '0'],
    ['Инвесторы', 'Нет', 'Нет'],
    ['Команда', 'Анонимная', '0 (участвуют наравне со всеми)'],
])
add_para(doc, 'В «Елене» все равны. Первые монеты появляются только через осаждение после запуска сети. Никаких привилегированных групп, никакого раннего доступа.')

# 9
add_heading(doc, '9. Сравнение с другими сетями', 1)
add_table(doc, ['Параметр', 'Биткоин', 'Ethereum', 'Елена'], [
    ['Эмиссия', 'Майнинг (PoW)', 'Стейкинг (PoS)', 'Осаждение'],
    ['Макс. предложение', '21 млн', 'Неограничено', '21 млн'],
    ['Власть', 'Хешрейт', 'Богатство', 'Репутация'],
    ['Комиссии', 'Добровольные', 'Газ (высокий)', '0.01% + сжигание'],
    ['Инфляция', 'Убывающая', 'Переменная', 'Убывающая'],
    ['Пре-майн', '0 (но Сатоши мог)', '72 млн (пресейл)', '0 абсолютно'],
    ['Справедливость', 'Майнеры богатеют', 'Богатые богатеют', 'Активные богатеют'],
    ['Gini коэф.', '0.88', '0.76', '0.52 (прогноз)'],
])

# 10
add_heading(doc, '10. Математическое резюме', 1)
math_block = '''╔════════════════════════════════════════════════════════════════╗
║                   ТОКЕНОМИКА В ФОРМУЛАХ                        ║
╠════════════════════════════════════════════════════════════════╣
║  МАКСИМАЛЬНОЕ ПРЕДЛОЖЕНИЕ                                      ║
║  M = 21_000_000 ELENA                                          ║
║  1 ELENA = 1_000_000 микро                                     ║
║  ЭМИССИЯ (ОСАЖДЕНИЕ)                                           ║
║  E(t) = E₀ × 2^(-t/2)  (халвинг каждые 2 года)                 ║
║  E₀ = 1_000_000 ELENA/год                                      ║
║  НАГРАДА УЗЛА                                                  ║
║  Rᵢ = E(t) × (Sᵢ × Tᵢ) / Σ(Sⱼ × Tⱼ) × (0.5 + 0.5 × rᵢ)        ║
║  КОМИССИЯ                                                      ║
║  C = 100 + (A × 1 / 10_000) × p  (в микро)                     ║
║  где A = amount в микро, p ∈ {1,2,10}                          ║
║  РАСПРЕДЕЛЕНИЕ КОМИССИИ                                        ║
║  C_storage = 0.5 × C   C_relay = 0.3 × C   C_burn = 0.2 × C    ║
║  РЕПУТАЦИЯ                                                     ║
║  r ∈ [0.01, 0.99]                                              ║
║  Δr = +0.0005 (relay) | +0.01 (alert) | -0.0001/день (decay)   ║
║  punish: r = 0.01 (double-spend)                               ║
║  СТЕЙКИНГ РЕПУТАЦИИ                                            ║
║  r_eff = r × (1 + 0.5 × s), s ∈ [0, 0.5]                       ║
║  штраф при нарушении: r × (1 + s)                               ║
╚════════════════════════════════════════════════════════════════╝'''
add_code(doc, math_block)

# 11
add_heading(doc, '11. Прогнозы и моделирование', 1)

add_para(doc, 'Сценарий 1: Ранний этап (1-й год)', bold=True)
add_para(doc, '• Активных узлов: 1 000')
add_para(doc, '• Транзакций/день: 10 000')
add_para(doc, '• Средняя репутация: 0.6')
add_para(doc, '• Эмиссия/день: ~2 700 ELENA')
add_para(doc, '• Комиссий собрано: ~100 ELENA/день')
add_para(doc, '• Сожжено: ~20 ELENA/день')

add_para(doc, 'Сценарий 2: Рост (2-й год)', bold=True)
add_para(doc, '• Активных узлов: 10 000')
add_para(doc, '• Транзакций/день: 100 000')
add_para(doc, '• Средняя репутация: 0.7')
add_para(doc, '• Эмиссия/день: ~2 000 ELENA (после халвинга)')
add_para(doc, '• Комиссий собрано: ~1 000 ELENA/день')
add_para(doc, '• Сожжено: ~200 ELENA/день')

add_para(doc, 'Сценарий 3: Зрелая сеть (5-й год)', bold=True)
add_para(doc, '• Активных узлов: 100 000')
add_para(doc, '• Транзакций/день: 1 000 000')
add_para(doc, '• Средняя репутация: 0.8')
add_para(doc, '• Эмиссия/день: ~500 ELENA')
add_para(doc, '• Комиссий собрано: ~10 000 ELENA/день')
add_para(doc, '• Сожжено: ~2 000 ELENA/день')

add_para(doc, 'Коэффициент Джини (распределение богатства):', bold=True)
add_formula(doc, 'Биткоин: 0.88  (очень неравномерно)')
add_formula(doc, 'Ethereum: 0.76  (лучше, но всё плохо)')
add_formula(doc, 'Традиционная экономика: 0.60–0.70')
add_formula(doc, 'Елена (прогноз): 0.52  (значительно равномернее)')

# 12
add_heading(doc, '12. Управление параметрами (Governance)', 1)
add_para(doc, 'Какие параметры можно менять голосованием:')
add_table(doc, ['Параметр', 'Ограничения', 'Периодичность'], [
    ['Базовая эмиссия (E₀)', '±50% от текущей', 'Не чаще 1 раза в год'],
    ['FEE_BASE', '10–1000 микро', 'Не чаще 1 раза в 3 мес'],
    ['FEE_RATE_BP', '0–10', 'Не чаще 1 раза в 3 мес'],
    ['Распределение комиссий', 'storage 40-60%, relay 20-40%, burn 10-30%', 'Не чаще 1 раза в 6 мес'],
    ['Пороги финальности', '0.8–0.99', 'В любой момент'],
])
add_para(doc, 'Какие параметры нельзя менять:')
add_para(doc, '❌ Максимальное предложение (21 млн)')
add_para(doc, '❌ Пост-квантовые алгоритмы (можно только апгрейдить)')
add_para(doc, '❌ Механизм коллизии якорей (ядро консенсуса)')
add_para(doc, '❌ Базовый диапазон репутации (0.01–0.99)')
add_para(doc, 'Голосование:', bold=True)
add_formula(doc, 'Вес голоса = reputation × (1 + 0.5 × stake)')
add_formula(doc, 'Для принятия: >66% голосов')
add_formula(doc, 'Для критических изменений: >80% + временная задержка 30 дней')

# Заключение
add_heading(doc, 'Заключение', 1)
add_para(doc, 'Токеномика «Елены» — это принципиально новая экономическая модель, где:')
add_para(doc, '✅ Власть = репутация, а не деньги')
add_para(doc, '✅ Эмиссия = вклад, а не майнинг')
add_para(doc, '✅ Комиссии = сжигаются, создавая дефляцию')
add_para(doc, '✅ 21 млн лимит — цифровая редкость')
add_para(doc, '✅ Абсолютно справедливый старт — ни у кого нет преимущества')
add_para(doc, 'Это экономика, в которой выгодно быть честным и активным, а не богатым и влиятельным.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('© 2026 Протокол Елена')
r.font.size = Pt(10)
r.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run('Лицензия: CC BY-SA 4.0')
r.font.size = Pt(10)
r.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out_path = '/Users/macbook/Desktop/TOKENOMICS_ELENA_v2.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
